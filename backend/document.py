import string
from docx import Document

def load_document(file_path_or_stream):
    """Load DOCX document."""
    return Document(file_path_or_stream)

def save_document(doc, file_path_or_stream):
    """Save DOCX document."""
    doc.save(file_path_or_stream)

def redact_string(text: str, mapping_dict: dict) -> str:
    """Redact string mapping."""
    if not text or not mapping_dict:
        return text

    matches = []
    
    for original, replacement in mapping_dict.items():
        if not original:
            continue
        
        start_idx = 0
        while True:
            idx = text.find(original, start_idx)
            if idx == -1:
                break
            
            span_start = idx
            span_end = idx + len(original)
            
            is_word = True
            if span_start > 0:
                char_before = text[span_start - 1]
                if char_before.isalnum():
                    is_word = False
            if span_end < len(text):
                char_after = text[span_end]
                if char_after.isalnum():
                    is_word = False

            if not is_word:
                start_idx = idx + 1
                continue

            overlap = False
            for m in matches:
                if span_start < m["end"] and span_end > m["start"]:
                    overlap = True
                    break
            
            if not overlap:
                matches.append({
                    "start": span_start,
                    "end": span_end,
                    "replacement": replacement
                })
            
            start_idx = idx + 1

    matches.sort(key=lambda x: x["start"], reverse=True)

    result = text
    for m in matches:
        result = result[:m["start"]] + m["replacement"] + result[m["end"]:]

    return result

def redact_paragraph(paragraph, mapping_dict):
    """Redact paragraph runs and text."""
    text = paragraph.text
    if not text:
        return

    stripped = text.strip()
    if not stripped:
        return

    if len(stripped) <= 2:
        return
    if stripped.isdigit():
        return
    if all(c in string.punctuation or c.isspace() for c in stripped):
        return
    
    lower_text = stripped.lower()
    if lower_text.startswith("page ") and lower_text[5:].strip().isdigit():
        return

    for run in paragraph.runs:
        if not run.text:
            continue
        run_stripped = run.text.strip()
        if not run_stripped or len(run_stripped) <= 2 or run_stripped.isdigit():
            continue
        run.text = redact_string(run.text, mapping_dict)

    full_text = paragraph.text
    needs_fallback = False
    for original in mapping_dict.keys():
        if original in full_text:
            needs_fallback = True
            break

    if needs_fallback:
        paragraph.text = redact_string(full_text, mapping_dict)

def redact_document(doc, mapping_dict):
    """Redact doc paragraphs and tables."""
    if not mapping_dict:
        return

    total_paragraphs = len(doc.paragraphs)
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx > 0 and idx % 100 == 0:
            print(f"Processing paragraph {idx}/{total_paragraphs}...")
        redact_paragraph(paragraph, mapping_dict)

    total_tables = len(doc.tables)
    print("Processing tables...")
    for idx, table in enumerate(doc.tables):
        if idx > 0 and idx % 10 == 0:
            print(f"Processing table {idx}/{total_tables}...")
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                redact_paragraph(paragraph, mapping_dict)
